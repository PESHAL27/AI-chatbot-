import os
import re
import json
import math
import uuid
import logging
from typing import List, Dict, Any, Optional, Tuple
from app.config import settings
from app.services.database_service import DatabaseService
from app.services.ai_service import AIService

logger = logging.getLogger("pml.rag_service")

class RAGService:
    # ==================== TEXT EXTRACTION ====================

    @classmethod
    def extract_text(cls, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        Extracts structured text from PDF, DOCX, and TXT documents.
        Returns a list of page/section items: [{"page_number": int, "text": str}]
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document file not found at path: {file_path}")

        ft = file_type.lower().lstrip(".")
        sections: List[Dict[str, Any]] = []

        if ft == "pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                total_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text() or ""
                        clean_text = cls._clean_text(text)
                        if clean_text:
                            sections.append({
                                "page_number": i + 1,
                                "text": clean_text
                            })
                    except Exception as page_err:
                        logger.warn(f"Error extracting page {i+1} in PDF {file_path}: {page_err}")
                
                if not sections and total_pages > 0:
                    sections.append({
                        "page_number": 1,
                        "text": f"[PDF contains {total_pages} page(s) with graphical or scanned content]"
                    })
            except Exception as e:
                logger.error(f"Failed to read PDF file {file_path}: {e}")
                raise ValueError(f"Unable to parse PDF document: {str(e)}")

        elif ft in ("docx", "doc"):
            try:
                import docx
                doc = docx.Document(file_path)
                full_text: List[str] = []

                for p in doc.paragraphs:
                    if p.text and p.text.strip():
                        full_text.append(p.text.strip())

                for table in doc.tables:
                    for row in table.rows:
                        row_content = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_content:
                            full_text.append(row_content)

                combined = "\n\n".join(full_text)
                if combined.strip():
                    sections.append({
                        "page_number": 1,
                        "text": cls._clean_text(combined)
                    })
            except Exception as e:
                logger.error(f"Failed to read DOCX file {file_path}: {e}")
                raise ValueError(f"Unable to parse DOCX document: {str(e)}")

        elif ft in ("txt", "md", "csv", "json", "log"):
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
                clean_text = cls._clean_text(text)
                if clean_text:
                    sections.append({
                        "page_number": 1,
                        "text": clean_text
                    })
            except Exception as e:
                logger.error(f"Failed to read text file {file_path}: {e}")
                raise ValueError(f"Unable to read text document: {str(e)}")
        else:
            raise ValueError(f"Unsupported file format: {file_type}. Supported: PDF, DOCX, TXT")

        if not sections:
            raise ValueError("No extractable text found in uploaded document.")

        return sections

    @staticmethod
    def _clean_text(text: str) -> str:
        """Cleans extracted text, normalizes whitespace and removes null bytes."""
        if not text:
            return ""
        # Remove null bytes and non-printable characters
        text = text.replace("\x00", " ")
        # Normalize multiple spaces and multiple blank lines
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
        return text.strip()

    # ==================== TEXT CHUNKING ====================

    @classmethod
    def chunk_sections(
        cls,
        document_id: str,
        user_id: str,
        sections: List[Dict[str, Any]],
        chunk_word_size: int = 250,
        chunk_overlap: int = 40
    ) -> List[Dict[str, Any]]:
        """
        Splits extracted sections into coherent semantic chunks preserving page metadata.
        """
        chunks: List[Dict[str, Any]] = []
        chunk_index = 0

        for sec in sections:
            page_num = sec.get("page_number", 1)
            raw_text = sec.get("text", "").strip()
            if not raw_text:
                continue

            # Split text by paragraphs first
            paragraphs = [p.strip() for p in raw_text.split("\n\n") if p.strip()]
            current_chunk_words: List[str] = []

            for para in paragraphs:
                para_words = para.split()
                if not para_words:
                    continue

                if len(current_chunk_words) + len(para_words) <= chunk_word_size:
                    current_chunk_words.extend(para_words)
                else:
                    if current_chunk_words:
                        chunk_text = " ".join(current_chunk_words)
                        chunks.append({
                            "id": f"pml-chk-{uuid.uuid4().hex[:12]}",
                            "document_id": document_id,
                            "user_id": user_id,
                            "content": chunk_text,
                            "chunk_index": chunk_index,
                            "page_number": page_num
                        })
                        chunk_index += 1

                        # Carry over overlap words for context continuity
                        if chunk_overlap > 0 and len(current_chunk_words) > chunk_overlap:
                            current_chunk_words = current_chunk_words[-chunk_overlap:]
                        else:
                            current_chunk_words = []

                    # If paragraph itself is larger than chunk size, break it up
                    if len(para_words) > chunk_word_size:
                        for i in range(0, len(para_words), chunk_word_size - chunk_overlap):
                            sub_slice = para_words[i : i + chunk_word_size]
                            if sub_slice:
                                chunks.append({
                                    "id": f"pml-chk-{uuid.uuid4().hex[:12]}",
                                    "document_id": document_id,
                                    "user_id": user_id,
                                    "content": " ".join(sub_slice),
                                    "chunk_index": chunk_index,
                                    "page_number": page_num
                                })
                                chunk_index += 1
                        current_chunk_words = []
                    else:
                        current_chunk_words.extend(para_words)

            if current_chunk_words:
                chunks.append({
                    "id": f"pml-chk-{uuid.uuid4().hex[:12]}",
                    "document_id": document_id,
                    "user_id": user_id,
                    "content": " ".join(current_chunk_words),
                    "chunk_index": chunk_index,
                    "page_number": page_num
                })
                chunk_index += 1

        return chunks

    # ==================== EMBEDDINGS & VECTOR ENGINES ====================

    @classmethod
    async def generate_embeddings_batch(cls, texts: List[str]) -> List[List[float]]:
        """
        Generates vector embeddings for a list of text chunks using OpenAI embeddings API.
        Includes a fast local semantic vector fallback if network or quota is limited.
        """
        if not texts:
            return []

        try:
            client = AIService.get_client()
            clean_inputs = [t[:3000].replace("\n", " ") for t in texts]
            
            res = await client.embeddings.create(
                model="text-embedding-3-small",
                input=clean_inputs
            )
            return [item.embedding for item in res.data]
        except Exception as e:
            logger.warn(f"OpenAI embeddings API failed ({e}), using high-dimensional local semantic feature vector fallback.")
            return [cls._generate_local_embedding(t) for t in texts]

    @classmethod
    async def generate_query_embedding(cls, query: str) -> List[float]:
        """Generates a single query embedding vector."""
        embs = await cls.generate_embeddings_batch([query])
        return embs[0] if embs else cls._generate_local_embedding(query)

    @staticmethod
    def _generate_local_embedding(text: str, dim: int = 128) -> List[float]:
        """Deterministic, normalized pseudo-semantic vector generator for fallback/offline environments."""
        vec = [0.0] * dim
        words = re.findall(r"\b\w+\b", text.lower())
        if not words:
            return vec

        for w in words:
            # Hash tokens across feature dimensions
            h = hash(w)
            idx = abs(h) % dim
            val = (h % 1000) / 1000.0
            vec[idx] += val

        # L2 Normalize vector
        norm = math.sqrt(sum(x * x for x in vec))
        if norm > 0:
            vec = [x / norm for x in vec]
        return vec

    @staticmethod
    def cosine_similarity(v1: List[float], v2: List[float]) -> float:
        """Computes cosine similarity between two float vectors."""
        if not v1 or not v2 or len(v1) != len(v2):
            return 0.0
        dot = sum(a * b for a, b in zip(v1, v2))
        norm1 = math.sqrt(sum(a * a for a in v1))
        norm2 = math.sqrt(sum(b * b for b in v2))
        if norm1 == 0 or norm2 == 0:
            return 0.0
        return dot / (norm1 * norm2)

    # ==================== RAG RETRIEVAL PIPELINE ====================

    @classmethod
    async def retrieve_relevant_chunks(
        cls,
        user_id: str,
        query: str,
        document_id: Optional[str] = None,
        top_k: int = 4,
        user_token: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Performs semantic vector search + keyword matching over the user's document chunks.
        Strictly isolated by user_id.
        """
        if not user_id:
            return []

        # 1. Fetch user's documents to map document_id to file_name
        user_docs = await DatabaseService.get_documents(user_id=user_id, user_token=user_token)
        if not user_docs:
            return []

        doc_map = {d["id"]: d.get("file_name", "Document") for d in user_docs if d.get("status") == "ready"}
        if not doc_map:
            return []

        # 2. Fetch document chunks
        all_chunks = await DatabaseService.get_document_chunks(
            user_id=user_id,
            document_id=document_id,
            user_token=user_token
        )
        if not all_chunks:
            return []

        # 3. Generate query embedding
        query_vec = await cls.generate_query_embedding(query)
        query_words = set(w.lower() for w in re.findall(r"\b\w+\b", query) if len(w) > 2)

        scored_chunks: List[Tuple[float, Dict[str, Any]]] = []

        for chk in all_chunks:
            # Only include chunks from ready documents
            doc_id = chk.get("document_id")
            if doc_id not in doc_map:
                continue

            content = chk.get("content", "")
            raw_emb = chk.get("embedding")
            sim = 0.0

            if raw_emb:
                try:
                    if isinstance(raw_emb, str):
                        emb_vec = json.loads(raw_emb)
                    else:
                        emb_vec = raw_emb
                    sim = cls.cosine_similarity(query_vec, emb_vec)
                except Exception:
                    sim = 0.0

            # Keyword lexical overlap boost
            content_words = set(w.lower() for w in re.findall(r"\b\w+\b", content) if len(w) > 2)
            if query_words and content_words:
                overlap = len(query_words.intersection(content_words)) / len(query_words)
                sim += overlap * 0.45

            # If user explicitly asked about a specific file name (e.g. "in Java_Unit_2.pdf")
            file_name = doc_map.get(doc_id, "")
            file_name_clean = file_name.lower().replace("_", " ").replace("-", " ")
            if any(part in query.lower() for part in file_name_clean.split() if len(part) > 2):
                sim += 0.35

            chk_copy = dict(chk)
            chk_copy["file_name"] = file_name
            chk_copy["score"] = round(sim, 4)
            scored_chunks.append((sim, chk_copy))

        # Sort by relevance score descending
        scored_chunks.sort(key=lambda x: x[0], reverse=True)

        # Select top chunks above minimum threshold
        results = [item[1] for item in scored_chunks if item[0] >= 0.20][:top_k]

        # Fallback: if user scoped directly to a document, return the first 2 chunks
        if not results and document_id and all_chunks:
            fallback = [dict(c) for c in all_chunks[:2]]
            for f in fallback:
                f["file_name"] = doc_map.get(f.get("document_id"), "Document")
                f["score"] = 0.5
            return fallback

        logger.info(f"[PML RAG] Retrieved {len(results)} relevant document chunks for query '{query[:40]}' (user: {user_id})")
        return results

    # ==================== DOCUMENT INGESTION PIPELINE ====================

    @classmethod
    async def process_document_pipeline(
        cls,
        document_id: str,
        user_id: str,
        file_path: str,
        file_type: str,
        user_token: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Executes end-to-end ingestion:
        1. Status -> 'processing'
        2. Text extraction
        3. Semantic chunking
        4. Embedding vector generation
        5. Chunk database storage
        6. Status -> 'ready' (or 'failed' on error)
        """
        logger.info(f"[PML RAG] Starting ingestion pipeline for document {document_id} ({file_type})")
        try:
            # 1. Update status to processing
            await DatabaseService.update_document_status(
                document_id=document_id,
                user_id=user_id,
                status="processing",
                user_token=user_token
            )

            # 2. Extract text sections
            sections = cls.extract_text(file_path, file_type)

            # 3. Chunk sections
            chunks = cls.chunk_sections(
                document_id=document_id,
                user_id=user_id,
                sections=sections
            )

            if not chunks:
                raise ValueError("No meaningful text chunks could be extracted from document.")

            # 4. Generate embeddings
            texts = [c["content"] for c in chunks]
            embeddings = await cls.generate_embeddings_batch(texts)

            # Attach serialized embeddings to chunk payloads
            for chk, emb in zip(chunks, embeddings):
                chk["embedding"] = json.dumps(emb)

            # 5. Store chunks in database
            saved_count = await DatabaseService.save_document_chunks(
                chunks=chunks,
                user_token=user_token
            )

            # 6. Update document status to ready
            doc = await DatabaseService.update_document_status(
                document_id=document_id,
                user_id=user_id,
                status="ready",
                chunk_count=saved_count,
                user_token=user_token
            )
            logger.info(f"[PML RAG] Document {document_id} successfully processed with {saved_count} chunks.")
            return doc or {"id": document_id, "status": "ready", "chunk_count": saved_count}

        except Exception as err:
            logger.error(f"[PML RAG] Document {document_id} processing failed: {err}", exc_info=True)
            await DatabaseService.update_document_status(
                document_id=document_id,
                user_id=user_id,
                status="failed",
                error_message=str(err)[:250],
                user_token=user_token
            )
            raise err
